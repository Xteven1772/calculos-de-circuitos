import dash
from dash import html, dcc, Input, Output, State, ctx
import dash_bootstrap_components as dbc
import plotly.graph_objs as go
import numpy as np
from docx import Document
import base64
import io
import pandas as pd
import os


app = dash.Dash(
    __name__,
    external_stylesheets=[dbc.themes.SLATE],
    title="Analizador BJT"
)

app.title = "Analizador BJT - Transistores"

# ----- Estilos personalizados -----
app.index_string = '''
<!DOCTYPE html>
<html>
    <head>
        {%metas%}
        <title>{%title%}</title>
        {%favicon%}
        {%css%}
        <style>
            body { background-color: #1e1e2f; font-family: 'Segoe UI', sans-serif; }
            .soft-box {
                background-color: rgba(255,255,255,0.05);
                padding: 15px;
                border-radius: 12px;
                box-shadow: 0 4px 30px rgba(0,0,0,0.1);
                backdrop-filter: blur(4px);
                border: 1px solid rgba(255,255,255,0.1);
                margin-bottom: 15px;
                transition: all 0.3s ease-in-out;
            }
            input:invalid {
                background-color: #ffcccc !important;
            }
            .error-input {
                background-color: #ffcccc !important;
            }
        </style>
    </head>
    <body>
        {%app_entry%}
        <footer>
            {%config%}
            {%scripts%}
            {%renderer%}
        </footer>
    </body>
</html>
'''

# ----- Funciones de utilidad -----

def interpretar_valor(valor_str):
    prefijos = {
        "E": 1e18, "P": 1e15, "T": 1e12, "G": 1e9, "M": 1e6, "k": 1e3,
        "h": 1e2, "da": 1e1, "d": 1e-1, "c": 1e-2, "m": 1e-3,
        "u": 1e-6, "µ": 1e-6, "n": 1e-9, "p": 1e-12, "f": 1e-15, "a": 1e-18
    }
    valor_str = valor_str.strip()
    if not valor_str:
        return None
    for p in sorted(prefijos, key=lambda x: -len(x)):
        if valor_str.lower().endswith(p.lower()):
            try:
                num = float(valor_str[:-len(p)])
                return num * prefijos[p]
            except:
                raise ValueError("Error al interpretar valor con prefijo")
    return float(valor_str)

def formatear_valor(valor):
    if abs(valor) >= 1:
        return f"{valor:.3f} A"
    elif abs(valor) >= 1e-3:
        return f"{valor*1e3:.3f} mA"
    elif abs(valor) >= 1e-6:
        return f"{valor*1e6:.3f} µA"
    else:
        return f"{valor*1e9:.3f} nA"

def exportar_a_word(resultados_dict):
    import io, base64
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    doc = Document()
    # Título
    doc.add_heading('Reporte de Análisis de Transistor BJT', 0)
    # Introducción
    doc.add_paragraph(
        "Este reporte presenta el análisis de un transistor BJT en distintas configuraciones. "
        "Se detallan los parámetros utilizados, los procesos de cálculo y los resultados obtenidos. "
        "Los valores deducidos o asumidos se indican con el símbolo '≅'."
    )
    # Procesos de cálculo
    def get_val(key):
        v = resultados_dict.get(key)
        if hasattr(v, 'children'):
            if isinstance(v.children, list):
                return ''.join(str(c) for c in v.children)
            else:
                return str(v.children)
        return str(v)

    procesos = [
        ("Cálculo de Ib", "Ib = (Vcc - Vbe) / (Rb + (β+1)·Re)",
         f"Ib = ({get_val('Vcc') if 'Vcc' in resultados_dict else '?'} - {get_val('Vbe') if 'Vbe' in resultados_dict else '?'}) / (" +
         f"{get_val('Rb') if 'Rb' in resultados_dict else '?'} + (" +
         f"{get_val('β') if 'β' in resultados_dict else '?'}+1)·{get_val('Re') if 'Re' in resultados_dict else '?'}) = {get_val('Ib')}") ,
        ("Cálculo de Ic", "Ic = β · Ib", f"Ic = {get_val('β')} · {get_val('Ib')} = {get_val('Ic')}") ,
        ("Cálculo de Ie", "Ie = Ic + Ib", f"Ie = {get_val('Ic')} + {get_val('Ib')} = {get_val('Ie')}") ,
        ("Cálculo de Ve", "Ve = Ie · Re", f"Ve = {get_val('Ie')} · {get_val('Re') if 'Re' in resultados_dict else '?'} = {get_val('Ve')}") ,
        ("Cálculo de Vb", "Vb = Ve + Vbe", f"Vb = {get_val('Ve')} + {get_val('Vbe') if 'Vbe' in resultados_dict else '?'} = {get_val('Vb')}") ,
        ("Cálculo de Vc", "Vc = Vcc - Ic · Rc (o Vcc si es colector común)", f"Vc = {get_val('Vc')}") ,
        ("Cálculo de Vce", "Vce = Vc - Ve", f"Vce = {get_val('Vc')} - {get_val('Ve')} = {get_val('Vce')}") ,
        ("Cálculo de Vbc", "Vbc = Vb - Vc", f"Vbc = {get_val('Vb')} - {get_val('Vc')} = {get_val('Vbc')}") ,
        ("Cálculo de Ic(sat)", "Ic(sat) = Vcc / Rc", f"Ic(sat) = {get_val('Vcc') if 'Vcc' in resultados_dict else '?'} / {get_val('Rc') if 'Rc' in resultados_dict else '?'} = {get_val('Ic(sat)')}") ,
        ("Cálculo de Pmax", "Pmax = Vce(sat) · Ic(sat)", f"Pmax = {get_val('Vce(sat)')} · {get_val('Ic(sat)')} = {get_val('Pmax')}")
    ]
    doc.add_heading('Procesos de cálculo', level=1)
    for nombre, formula, desarrollo in procesos:
        doc.add_paragraph(nombre, style='List Bullet')
        p = doc.add_paragraph()
        p.add_run(formula + "\n").bold = True
        p.add_run(desarrollo)
    # Tabla de resultados
    doc.add_heading('Resultados', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parámetro'
    hdr_cells[1].text = 'Valor'
    for clave, valor in resultados_dict.items():
        if hasattr(valor, 'children'):
            if isinstance(valor.children, list):
                val = ''.join(str(c) for c in valor.children)
            else:
                val = str(valor.children)
        else:
            val = str(valor)
        row_cells = table.add_row().cells
        row_cells[0].text = clave
        row_cells[1].text = val
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    b64 = base64.b64encode(buffer.read()).decode('utf-8')
    return f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}"

# ----- Lógica principal -----

def calcular_y_graficar(config, Vcc, Rc, Rb, Re, beta, Vbe):
    # (Se mueve la creación de valores_efectivos más abajo, después de definir *_val)

    # Valores típicos
    valores_defecto = {
        "Vcc": 12.0,   # Voltios
        "Rc": 1000.0,  # Ohms
        "Rb": 100000.0, # Ohms
        "Re": 1000.0,  # Ohms
        "β": 100.0,    # Sin unidad
        "Vbe": 0.7     # Voltios
    }

    deducidos = []
    usados_por_defecto = []

    # Intentar deducir cada parámetro si es posible
    # 1. Vcc
    try:
        Vcc_val = interpretar_valor(Vcc)
        if Vcc_val is None:
            raise ValueError()
        Vcc_src = 'directo'
    except:
        Vcc_val = None
        Vcc_src = None
    # 2. Rc
    try:
        Rc_val = interpretar_valor(Rc)
        if Rc_val is None:
            raise ValueError()
        Rc_src = 'directo'
    except:
        Rc_val = None
        Rc_src = None
    # 3. Rb
    try:
        Rb_val = interpretar_valor(Rb)
        if Rb_val is None:
            raise ValueError()
        Rb_src = 'directo'
    except:
        Rb_val = None
        Rb_src = None
    # 4. Re
    try:
        Re_val = interpretar_valor(Re)
        if Re_val is None:
            raise ValueError()
        Re_src = 'directo'
    except:
        Re_val = None
        Re_src = None
    # 5. β
    try:
        beta_val = float(beta)
        if beta_val is None:
            raise ValueError()
        beta_src = 'directo'
    except:
        beta_val = None
        beta_src = None
    # 6. Vbe
    try:
        Vbe_val = interpretar_valor(Vbe)
        if Vbe_val is None:
            raise ValueError()
        Vbe_src = 'directo'
    except:
        Vbe_val = None
        Vbe_src = None

    # Deducción básica para parámetros faltantes
    # Solo deducir si hay datos suficientes
    # Deducir β si Ic e Ib están disponibles
    Ic_deducido = None
    Ib_deducido = None
    Ie_deducido = None
    # Primero, intentar deducir β
    if beta_val is None:
        # Si Ic e Ib están disponibles, deducir β
        try:
            Ic_test = float(interpretar_valor(Rc))  # No hay campo Ic directo, así que no se puede deducir
        except:
            Ic_test = None
        try:
            Ib_test = float(interpretar_valor(Rb))  # No hay campo Ib directo, así que no se puede deducir
        except:
            Ib_test = None
        # No hay campo Ic ni Ib directo, así que no se puede deducir β
        pass
    # Deducir Ib si Ic y β están disponibles
    if (Rb_val is None) and (beta_val is not None):
        # Si el usuario ingresó Ic (no hay campo), no se puede deducir Ib
        pass
    # Deducir Ic si Ib y β están disponibles
    if (Rc_val is None) and (beta_val is not None):
        # Si el usuario ingresó Ib (no hay campo), no se puede deducir Ic
        pass
    # Deducir Vcc si Rc, Ic y Vc están disponibles
    # No hay campo Vc ni Ic directo
    # Deducir Rc si Vcc, Ic y Vc están disponibles
    # No hay campo Vc ni Ic directo
    # Deducir Vbe si Vb y Ve están disponibles
    # No hay campo Vb ni Ve directo

    # Si no se pudo deducir, usar valor típico
    if Vcc_val is None:
        Vcc_val = valores_defecto["Vcc"]
        usados_por_defecto.append("Vcc")
    if Rc_val is None:
        Rc_val = valores_defecto["Rc"]
        usados_por_defecto.append("Rc")
    if Rb_val is None:
        Rb_val = valores_defecto["Rb"]
        usados_por_defecto.append("Rb")
    if Re_val is None:
        Re_val = valores_defecto["Re"]
        usados_por_defecto.append("Re")
    if beta_val is None:
        beta_val = valores_defecto["β"]
        usados_por_defecto.append("β")
    if Vbe_val is None:
        Vbe_val = valores_defecto["Vbe"]
        usados_por_defecto.append("Vbe")

    # Ahora sí, guardar los valores efectivos usados para historial y curvas dinámicas
    valores_efectivos = {
        "Vcc": Vcc_val,
        "Rc": Rc_val,
        "Rb": Rb_val,
        "Re": Re_val,
        "β": beta_val,
        "Vbe": Vbe_val
    }

    # Realizar cálculo con los valores (originales, deducidos o por defecto)
    Ib = Ic = Ie = Vb = Ve = Vc = Vce = Vbc = 0

    if config == "Emisor común":
        divisor = Rb_val + (beta_val + 1) * Re_val if (Rb_val or Re_val) else 1
        Ib = (Vcc_val - Vbe_val) / divisor
        Ic = beta_val * Ib
        Ie = Ic + Ib

    elif config == "Base común":
        Ie = (Vcc_val - Vbe_val) / (Re_val + Rc_val)
        Ic = (beta_val / (beta_val + 1)) * Ie
        Ib = Ie - Ic

    elif config == "Colector común":
        divisor = Rb_val + (beta_val + 1) * Re_val if (Rb_val or Re_val) else 1
        Ib = (Vcc_val - Vbe_val) / divisor
        Ie = (beta_val + 1) * Ib
        Ic = beta_val * Ib

    Ve = Ie * Re_val
    Vb = Ve + Vbe_val
    Vc = Vcc_val if config == "Colector común" else Vcc_val - Ic * Rc_val
    Vce = Vc - Ve
    Vbc = Vb - Vc

    Vce_sat = 0.2
    Ic_sat = Vcc_val / Rc_val if Rc_val else 0
    Pmax = Vce_sat * Ic_sat

    estado = "SATURACIÓN" if Vce < 0.2 else "ACTIVA" if Ic > 0 else "CORTE"


    # Badge visual para el estado
    estado_color = {
        "ACTIVA": "success",
        "SATURACIÓN": "warning",
        "CORTE": "danger"
    }
    # Emoji visual para el estado
    emoji_estado = {
        "ACTIVA": "🟢",
        "SATURACIÓN": "🟡",
        "CORTE": "🔴"
    }
    badge_estado = html.Span([
        estado,
        html.Span(f" {emoji_estado.get(estado, '')}", style={"fontSize": "1.2em"})
    ], className=f"badge bg-{estado_color.get(estado, 'secondary')} mx-2", style={"fontSize": "1em"})

    # Mostrar ≅ para los parámetros deducidos o asumidos, con fondo amarillo suave e ícono
    def approx(val, nombre):
        if nombre in usados_por_defecto:
            # Siempre mostrar la unidad usando formatear_valor
            return html.Span([
                html.Span("⚠️", style={"marginRight": "3px", "fontSize": "1em"}),
                f"≅ {formatear_valor(val)}"
            ], style={"background": "#fff3cd", "color": "#856404", "padding": "2px 6px", "borderRadius": "6px", "fontWeight": "bold"})
        return formatear_valor(val) if nombre not in ['Vcc','Rc','Rb','Re','β','Vbe'] else val

    resultados_dict = {
        "Estado del transistor": badge_estado,
        "Ib": approx(Ib, "Rb"),
        "Ic": approx(Ic, "Rc"),
        "Ie": approx(Ie, "Re"),
        "Vb": f"{Vb:.2f} V",
        "Ve": f"{Ve:.2f} V",
        "Vc": f"{Vc:.2f} V",
        "Vce": f"{Vce:.2f} V",
        "Vbc": f"{Vbc:.2f} V",
        "Ic(sat)": approx(Ic_sat, "Rc"),
        "Vce(sat)": f"{Vce_sat:.2f} V",
        "Pmax": f"{Pmax:.3f} W"
    }

    # Aviso de deducidos y asumidos
    aviso = None
    if usados_por_defecto or deducidos:
        items = []
        if deducidos:
            items += [html.Li([
                html.Span("🧮", style={"marginRight": "4px"}),
                f"{p} deducido"
            ]) for p in deducidos]
        if usados_por_defecto:
            items += [html.Li([
                html.Span("⚠️", style={"marginRight": "4px"}),
                f"{p} ≅ {valores_defecto[p] if p != 'β' else int(valores_defecto[p])}"
            ]) for p in usados_por_defecto]
        aviso = html.Div([
            html.B("Parámetros deducidos o asumidos automáticamente:"),
            html.Ul(items, style={"color": "#856404"})
        ], style={"marginBottom": "10px", "background": "#fff3cd", "borderRadius": "8px", "padding": "8px 12px"})

    resultados = html.Div([
        aviso if aviso else None,
        html.Table([
            html.Thead(html.Tr([html.Th("Parámetro"), html.Th("Valor")])),
            html.Tbody([
                html.Tr([
                    html.Td(k),
                    html.Td(v)
                ]) for k, v in resultados_dict.items()
            ])
        ], className="table table-dark table-striped soft-box")
    ])

    fig = go.Figure()
    fig.add_trace(go.Scatter(x=[0, Vcc_val], y=[Ic_sat, 0], mode='lines', name='Recta de carga'))
    fig.add_trace(go.Scatter(x=[Vce], y=[Ic], mode='markers', name='Punto Q', marker=dict(size=10, color='red')))
    fig.update_layout(title="Recta de carga y punto Q", xaxis_title="VCE (V)", yaxis_title="IC (A)", template="plotly_dark")

    return resultados, fig, resultados_dict, valores_efectivos

    Ib = Ic = Ie = Vb = Ve = Vc = Vce = Vbc = 0

    if config == "Emisor común":
        divisor = Rb + (beta + 1) * Re if (Rb or Re) else 1
        Ib = (Vcc - Vbe) / divisor
        Ic = beta * Ib
        Ie = Ic + Ib

    elif config == "Base común":
        Ie = (Vcc - Vbe) / (Re + Rc)
        Ic = (beta / (beta + 1)) * Ie
        Ib = Ie - Ic

    elif config == "Colector común":
        divisor = Rb + (beta + 1) * Re if (Rb or Re) else 1
        Ib = (Vcc - Vbe) / divisor
        Ie = (beta + 1) * Ib
        Ic = beta * Ib

    Ve = Ie * Re
    Vb = Ve + Vbe
    Vc = Vcc if config == "Colector común" else Vcc - Ic * Rc
    Vce = Vc - Ve
    Vbc = Vb - Vc

    Vce_sat = 0.2
    Ic_sat = Vcc / Rc if Rc else 0
    Pmax = Vce_sat * Ic_sat

    estado = "SATURACIÓN" if Vce < 0.2 else "ACTIVA" if Ic > 0 else "CORTE"

    resultados_dict = {
        "Estado del transistor": estado,
        "Ib": formatear_valor(Ib),
        "Ic": formatear_valor(Ic),
        "Ie": formatear_valor(Ie),
        "Vb": f"{Vb:.2f} V",
        "Ve": f"{Ve:.2f} V",
        "Vc": f"{Vc:.2f} V",
        "Vce": f"{Vce:.2f} V",
        "Vbc": f"{Vbc:.2f} V",
        "Ic(sat)": formatear_valor(Ic_sat),
        "Vce(sat)": f"{Vce_sat:.2f} V",
        "Pmax": f"{Pmax:.3f} W"
    }

    resultados = html.Table([
        html.Thead(html.Tr([html.Th("Parámetro"), html.Th("Valor")])),
        html.Tbody([
            html.Tr([html.Td(k), html.Td(v)]) for k, v in resultados_dict.items()
        ])
    ], className="table table-dark table-striped soft-box")

    fig = go.Figure()
    fig.add_trace(go.Scatter(x=[0, Vcc], y=[Ic_sat, 0], mode='lines', name='Recta de carga'))
    fig.add_trace(go.Scatter(x=[Vce], y=[Ic], mode='markers', name='Punto Q', marker=dict(size=10, color='red')))
    fig.update_layout(title="Recta de carga y punto Q", xaxis_title="VCE (V)", yaxis_title="IC (A)", template="plotly_dark")

    return resultados, fig, resultados_dict

# ----- Diseño principal -----

import json
import os
HISTORIAL_PATH = "historial.json"
def cargar_historial():
    if os.path.exists(HISTORIAL_PATH):
        try:
            with open(HISTORIAL_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return []
    return []

def guardar_historial(historial):
    try:
        with open(HISTORIAL_PATH, "w", encoding="utf-8") as f:
            json.dump(historial, f, ensure_ascii=False, indent=2)
    except:
        pass

historial = cargar_historial()

def ayuda_parametro(param):
    ayudas = {
        "Vcc": "Voltaje de alimentación. Ejemplo: 12V, 24V.",
        "Rc": "Resistencia de colector. Ejemplo: 1k, 2.2k.",
        "Rb": "Resistencia de base. Ejemplo: 100k, 220k.",
        "Re": "Resistencia de emisor. Ejemplo: 1k, 470.",
        "β": "Ganancia de corriente (beta). Ejemplo: 100.",
        "Vbe": "Voltaje base-emisor. Típico: 0.7V."
    }
    return ayudas.get(param, "")

def input_with_help(campo):
    # Usar un ícono Unicode más claro y dbc.Tooltip para el tooltip
    info_id = f"info-{campo}"
    return html.Div([
        dbc.Label([
            campo,
            html.Span(" \u2139\ufe0f", id=info_id, style={"cursor": "pointer", "color": "#17a2b8", "fontWeight": "bold", "marginLeft": "6px"})
        ]),
        dbc.Tooltip(
            ayuda_parametro(campo),
            target=info_id,
            placement="right",
            style={"fontSize": "0.95em"}
        ),
        dbc.Input(id=campo, placeholder=campo, type="text", className="mb-2", value="")
    ])

app.layout = dbc.Container([
    html.H2("Analizador de Transistor BJT", className="my-3 text-center text-light"),

    dbc.Row([
        dbc.Col([
            html.Div([
                dbc.Label("Configuración"),
                dcc.Dropdown(
                    id="config",
                    options=[
                        {"label": "Emisor común", "value": "Emisor común"},
                        {"label": "Base común", "value": "Base común"},
                        {"label": "Colector común", "value": "Colector común"}
                    ],
                    value="Emisor común",
                    className="mb-3"
                ),
                *[input_with_help(campo) for campo in ["Vcc", "Rc", "Rb", "Re", "β", "Vbe"]],
                dbc.Button("Calcular", id="btn-calc", className="btn btn-success mt-2"),
                html.Br(),
                html.A("Descargar Word", id="descarga-word", href="", download="resultado.docx", target="_blank", className="btn btn-secondary mt-2", style={"display": "none"}),
                html.Br(),
                html.A("Descargar PDF", id="descarga-pdf", href="", download="resultado.pdf", target="_blank", className="btn btn-secondary mt-2", style={"display": "none"}),
                html.Div(id="validacion-campos", className="mt-2")
            ], className="soft-box")
        ], md=4),

        dbc.Col([
            dcc.Tabs(id="tabs", value="tab1", children=[
                dcc.Tab(label='Resultados', value='tab1'),
                dcc.Tab(label='Gráfica', value='tab2'),
                dcc.Tab(label='Curvas Dinámicas', value='tab3'),
                dcc.Tab(label='Historial', value='tab4')
            ]),
            html.Div(id="contenido_tab")
        ], md=8)
    ])
], fluid=True)

@app.callback(
    Output("contenido_tab", "children"),
    Output("descarga-word", "href"),
    Output("descarga-word", "style"),
    Output("descarga-pdf", "href"),
    Output("descarga-pdf", "style"),
    Output("validacion-campos", "children"),
    Input("tabs", "value"),
    Input("btn-calc", "n_clicks"),
    Input("config", "value"),
    Input("Vcc", "value"), Input("Rc", "value"), Input("Rb", "value"),
    Input("Re", "value"), Input("β", "value"), Input("Vbe", "value")
)
def actualizar_tabs(tab, n, config, Vcc, Rc, Rb, Re, beta, Vbe):
    # Validación visual y de rango

    campos = {"Vcc": Vcc, "Rc": Rc, "Rb": Rb, "Re": Re, "β": beta, "Vbe": Vbe}
    errores = []
    for k, v in campos.items():
        if v is not None and v != "":
            try:
                val = interpretar_valor(v)
                if k == "Vcc" and not (1 <= val <= 100):
                    errores.append("Vcc fuera de rango (1-100V)")
                if k == "β" and not (20 <= float(v) <= 500):
                    errores.append("β fuera de rango (20-500)")
            except:
                errores.append(f"{k} inválido")
    validacion = html.Ul([html.Li(e, style={"color": "#ff5555"}) for e in errores]) if errores else None

    global historial
    resultados, grafico, resultados_dict, valores_efectivos = calcular_y_graficar(config, Vcc, Rc, Rb, Re, beta, Vbe)
    href_word = exportar_a_word(resultados_dict)

    # Guardar historial solo si se presiona Calcular y no hay errores
    if ctx.triggered_id == "btn-calc" and not errores:
        hist = {"config": config}
        for k in ["Vcc", "Rc", "Rb", "Re", "β", "Vbe"]:
            hist[k] = valores_efectivos[k]
        historial.append(hist)
        guardar_historial(historial)

    # Tabla de ejemplos
    tabla_ejemplo = html.Table([
        html.Thead(html.Tr([html.Th("Parámetro"), html.Th("Ejemplo")])) ,
        html.Tbody([
            html.Tr([html.Td("Vcc"), html.Td("12 V")]),
            html.Tr([html.Td("Rc"), html.Td("1 kΩ")]),
            html.Tr([html.Td("Rb"), html.Td("100 kΩ")]),
            html.Tr([html.Td("Re"), html.Td("1 kΩ")]),
            html.Tr([html.Td("β"), html.Td("100")]),
            html.Tr([html.Td("Vbe"), html.Td("0.7 V")]),
        ])
    ], className="table table-bordered table-info soft-box")



    # Mostrar resultados en tiempo real: si hay errores, no mostrar resultados
    if any(errores):
        return validacion, "", {"display": "none"}, "", {"display": "none"}, validacion

    # Si no se ha hecho cálculo y no hay cambios, limpiar todo
    if not n and ctx.triggered_id != "btn-calc":
        return "", "", {"display": "none"}, "", {"display": "none"}, None


    # Si hay valores faltantes, mostrar tabla de faltantes en el área de resultados (derecha)

    # Detectar si el resultado es un mensaje de faltantes
    if isinstance(resultados, html.Div) and "Valores faltantes" in str(resultados):
        faltantes = str(resultados.children).split(": ")[-1].replace("⚠️ Valores faltantes: ", "").replace(")',)", "").replace("')", "").split(", ")
        tabla_faltantes = html.Table([
            html.Thead(html.Tr([html.Th("Parámetro faltante")])) ,
            html.Tbody([
                html.Tr([html.Td(f)]) for f in faltantes if f and f != "{}"
            ])
        ], className="table table-bordered table-warning soft-box")
        return tabla_faltantes, "", {"display": "none"}, "", {"display": "none"}, None


    if tab == "tab1":
        # Mejora visual: caja con sombra, separación, íconos
        return html.Div([
            html.Div([
                html.H4("Resultados del análisis", style={"color": "#00bfff", "marginBottom": "10px"}),
                resultados
            ], className="soft-box", style={"background": "rgba(0,191,255,0.07)", "border": "1.5px solid #00bfff", "boxShadow": "0 2px 12px #00bfff33"})
        ]), href_word, {"display": "inline-block"}, "", {"display": "none"}, None
    elif tab == "tab2":
        return dcc.Graph(figure=grafico), href_word, {"display": "inline-block"}, "", {"display": "none"}, None
    elif tab == "tab3":
        try:
            Vce_range = np.linspace(0, valores_efectivos["Vcc"], 100)
            divisor = valores_efectivos["Rb"] + (valores_efectivos["β"]+1)*valores_efectivos["Re"]
            Ib = (valores_efectivos["Vcc"] - valores_efectivos["Vbe"]) / divisor
            Ic_curva = [valores_efectivos["β"] * Ib for _ in Vce_range]
            fig_curvas = go.Figure()
            fig_curvas.add_trace(go.Scatter(x=Vce_range, y=Ic_curva, mode='lines', name='Curva IC vs VCE'))
            fig_curvas.update_layout(title="Curva Dinámica IC vs VCE", xaxis_title="VCE (V)", yaxis_title="IC (A)", template="plotly_dark")
            return dcc.Graph(figure=fig_curvas), href_word, {"display": "inline-block"}, "", {"display": "none"}, None
        except:
            return html.Div("Error al calcular curva dinámica. Revisa los valores."), href_word, {"display": "inline-block"}, "", {"display": "none"}, None
    elif tab == "tab4":
        if not historial:
            return html.Div("No hay cálculos previos."), href_word, {"display": "inline-block"}, "", {"display": "none"}, None
        # Mejora visual: tabla con colores y separación
        tabla_hist = html.Table([
            html.Thead(html.Tr([
                html.Th("Config"), html.Th("Vcc"), html.Th("Rc"), html.Th("Rb"), html.Th("Re"), html.Th("β"), html.Th("Vbe")
            ])),
            html.Tbody([
                html.Tr([
                    html.Td(h["config"]), html.Td(h["Vcc"]), html.Td(h["Rc"]), html.Td(h["Rb"]), html.Td(h["Re"]), html.Td(h["β"]), html.Td(h["Vbe"])
                ], style={"background": "#23233a" if i%2 else "#1e1e2f"}) for i, h in enumerate(historial)
            ])
        ], className="table table-bordered table-info soft-box", style={"boxShadow": "0 2px 8px #00bfff33", "marginTop": "10px"})
        return html.Div([
            html.H5("Historial de cálculos", style={"color": "#00bfff"}),
            tabla_hist
        ]), href_word, {"display": "inline-block"}, "", {"display": "none"}, None

    # (Eliminado retorno extra de tabla_ejemplo para mantener 6 valores)
    # Si quieres mostrar la tabla de ejemplo, inclúyela en el campo de validación o en contenido_tab, pero no como valor extra.

    resultados, grafico, resultados_dict = calcular_y_graficar(config, Vcc, Rc, Rb, Re, beta, Vbe)
    href_word = exportar_a_word(resultados_dict)

    if tab == "tab1":
        return resultados, href_word, {"display": "inline-block"}, None
    elif tab == "tab2":
        return dcc.Graph(figure=grafico), href_word, {"display": "inline-block"}, None
    elif tab == "tab3":
        try:
            Vce_range = np.linspace(0, interpretar_valor(Vcc), 100)
            divisor = interpretar_valor(Rb) + (float(beta)+1)*interpretar_valor(Re)
            Ib = (interpretar_valor(Vcc) - interpretar_valor(Vbe)) / divisor
            Ic_curva = [float(beta) * Ib for _ in Vce_range]
            fig_curvas = go.Figure()
            fig_curvas.add_trace(go.Scatter(x=Vce_range, y=Ic_curva, mode='lines', name='Curva IC vs VCE'))
            fig_curvas.update_layout(title="Curva Dinámica IC vs VCE", xaxis_title="VCE (V)", yaxis_title="IC (A)", template="plotly_dark")
            return dcc.Graph(figure=fig_curvas), href_word, {"display": "inline-block"}, None
        except:
            return html.Div("Error al calcular curva dinámica. Revisa los valores."), href_word, {"display": "inline-block"}, None
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8050))  # Usa el puerto de Render o 8050 por defecto
    app.run(host="0.0.0.0", port=port, debug=True)

